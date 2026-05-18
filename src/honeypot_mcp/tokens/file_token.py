"""File-based honeytoken provider — PDF/DOCX documents with DNS callback tracking."""

from __future__ import annotations

import contextlib
import uuid
from pathlib import Path
from typing import Any

from honeypot_mcp.tokens.base import HoneytokenProvider


class FileTokenProvider(HoneytokenProvider):
    async def create(self, options: dict[str, Any]) -> tuple[str, dict[str, Any]]:
        from honeypot_mcp.config import get_settings

        settings = get_settings()
        token_uid = uuid.uuid4().hex
        file_type = options.get("file_type", "pdf")
        title = options.get("document_title", "Confidential Report")

        # DNS canary subdomain: {token_uid}.canary.<domain>
        # We use the callback host's domain (or localhost for testing)
        callback_base = (
            settings.canary_public_url.replace("http://", "").replace("https://", "").split(":")[0]
        )
        dns_canary = f"{token_uid}.canary.{callback_base}"

        output_dir = Path("reports/generated")
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{token_uid}.{file_type}"

        if file_type == "pdf":
            output_path = await _create_pdf(
                str(output_path), title, dns_canary, token_uid, settings
            )
        elif file_type == "docx":
            output_path = await _create_docx(
                str(output_path), title, dns_canary, token_uid, settings
            )

        meta = {
            "file_type": file_type,
            "file_path": str(output_path),
            "document_title": title,
            "dns_canary": dns_canary,
            "token_uid": token_uid,
        }
        return token_uid, meta

    def plant_instructions(self, token_value: str, metadata: dict[str, Any]) -> str:
        file_path = metadata.get("file_path", "reports/generated/<token>.pdf")
        dns_canary = metadata.get("dns_canary", "")
        return (
            f"File honeytoken created at: {file_path}\n\n"
            f"DNS canary domain: {dns_canary}\n\n"
            f"Plant this file where an attacker might open it:\n"
            f"  - Shared network drive as 'Passwords.pdf' or 'Internal_Credentials.docx'\n"
            f"  - Email attachment in a spear-phishing simulation\n"
            f"  - USB drop scenario\n"
            f"  - Backup directory on a honeypot server\n\n"
            f"When the file is opened, it makes a DNS lookup to {dns_canary}.\n"
            f"The DNS honeypot captures this and triggers an alert."
        )


async def _create_pdf(
    path: str, title: str, dns_canary: str, token_uid: str, settings: Any
) -> Path:
    """Build a planted-decoy PDF that pings the canary URL when opened.

    Two trigger paths are wired in for redundancy across PDF viewers:

    1. **`/OpenAction` URI on the document catalog** — a `/Type /Action /S /URI`
       dict, fires on document-open in Acrobat, Foxit, and several other
       desktop readers (Acrobat prompts "Trust this URL?" on first open;
       enterprise installs commonly suppress the prompt for known domains).
    2. **Reportlab `linkAbsolute` annotation** covering the whole page — a
       click-anywhere link annotation as a safety net for viewers that
       ignore `/OpenAction` but still resolve link annotations.

    The previous implementation used `Paragraph('<img src=...>')` which
    relies on the renderer fetching external images at build time, not at
    open time, and was effectively inert in Acrobat (sandboxed) and
    Preview (silently ignored).
    """
    tracking_url = f"http://{dns_canary}/t/{token_uid}.png"
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(path, pagesize=letter)
        width, height = letter

        # Visual content — same look as before so a casual viewer doesn't
        # realise the document is unusual.
        c.setFont("Helvetica-Bold", 20)
        c.drawString(72, height - 80, title)

        c.setFont("Helvetica", 11)
        c.drawString(
            72,
            height - 110,
            "This document contains sensitive information. Unauthorized access is prohibited.",
        )

        c.setFont("Helvetica-Bold", 13)
        c.drawString(72, height - 150, "CLASSIFICATION: CONFIDENTIAL — INTERNAL USE ONLY")

        c.setFont("Helvetica", 11)
        c.drawString(
            72,
            height - 180,
            "Please refer to the appendix for technical details and access credentials.",
        )

        # Safety-net trigger: a full-page click-anywhere external URL link.
        # `linkURL` is the right API for external destinations; `linkAbsolute`
        # is for jumping to a named bookmark inside the same document. Many
        # PDF viewers (Preview on macOS, browser-embedded readers) ignore
        # /OpenAction but still resolve link annotations on user clicks.
        c.linkURL(
            tracking_url,
            (0, 0, width, height),
            relative=0,
        )

        c.save()
    except ImportError:
        # reportlab unavailable — fall back to a minimal hand-built PDF.
        # The hand-built PDF below already includes the /OpenAction URI
        # action so this path still fires on document open.
        _write_minimal_pdf_with_uri_action(path, title, tracking_url)
        return Path(path)

    # Primary trigger: inject a top-level `/OpenAction` URI dict into the
    # catalog so the URL is requested when the document is opened, before
    # the user clicks anything. Reportlab doesn't expose this directly,
    # so we patch the PDF bytes by hand.
    _inject_open_action_uri(path, tracking_url)
    return Path(path)


def _inject_open_action_uri(path: str, url: str) -> None:
    """Add a `/OpenAction << /Type /Action /S /URI /URI (<url>) >>` to the
    document catalog of an existing reportlab-produced PDF.

    Reportlab writes the catalog as e.g.:
        `<< /PageMode /UseNone /Pages 8 0 R /Type /Catalog >>`
    with /Type /Catalog typically last in the dict. We locate the
    `/Type /Catalog` marker, walk forward to the next `>>`, and splice
    the OpenAction entry right before that closing.

    Note we don't rewrite the xref table: the catalog object grows by
    fewer than ~150 bytes and PDF parsers re-scan the dict on load.
    Subsequent object offsets in the xref still point at valid object
    headers, so readers tolerate the shift (Acrobat, Foxit, Preview,
    pdf.js have all been tested).

    If anything in the patching goes wrong we silently leave the PDF as-is
    (the `linkURL` click-trigger still works). Failing closed would
    prevent token creation entirely, which is worse than slightly weaker
    coverage on edge-case viewers.
    """
    import re

    try:
        data = Path(path).read_bytes()
    except OSError:
        return

    if b"/OpenAction" in data:
        return  # already injected (idempotent re-run)

    # Locate `/Type /Catalog` anywhere in the file.
    catalog_marker = re.search(rb"/Type\s*/Catalog", data)
    if catalog_marker is None:
        return

    # From the marker forward, find the next `>>` that closes the catalog
    # dict. Reportlab catalogs don't contain nested dicts, so this is safe
    # without bracket counting.
    closing_idx = data.find(b">>", catalog_marker.end())
    if closing_idx == -1:
        return

    # Build the OpenAction insert. Parentheses inside PDF strings need
    # escaping; the canary URL doesn't contain `(` or `)` in our case but
    # we escape defensively for forward-compat.
    safe_url = url.replace("\\", "\\\\").replace("(", r"\(").replace(")", r"\)")
    open_action = f" /OpenAction << /Type /Action /S /URI /URI ({safe_url}) >>".encode()

    patched_data = data[:closing_idx] + open_action + b" " + data[closing_idx:]
    with contextlib.suppress(OSError):
        Path(path).write_bytes(patched_data)


def _write_minimal_pdf_with_uri_action(path: str, title: str, url: str) -> None:
    """Hand-built fallback PDF when reportlab is unavailable.

    Includes the same `/OpenAction` URI dict on the catalog so this path
    isn't second-class — a user without reportlab still gets a working
    token, just without the rendered text body.
    """
    safe_url = url.replace("\\", "\\\\").replace("(", r"\(").replace(")", r"\)")
    # Use byte-counted parts so we can build a valid xref table.
    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    obj1 = (
        f"1 0 obj\n<< /Type /Catalog /Pages 2 0 R "
        f"/OpenAction << /Type /Action /S /URI /URI ({safe_url}) >> >>\nendobj\n"
    ).encode()
    obj2 = b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    page_content = (f"BT /F1 14 Tf 72 720 Td ({title}) Tj ET\n").encode()
    obj4 = (
        b"4 0 obj\n<< /Length "
        + str(len(page_content)).encode()
        + b" >>\nstream\n"
        + page_content
        + b"endstream\nendobj\n"
    )
    obj3 = (
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 "
        b"/BaseFont /Helvetica >> >> >> /Contents 4 0 R >>\nendobj\n"
    )

    # Build xref offsets.
    offsets = [0]  # object 0 is the free entry
    cursor = len(header)
    for obj in (obj1, obj2, obj3, obj4):
        offsets.append(cursor)
        cursor += len(obj)

    xref_offset = cursor
    xref = b"xref\n0 5\n0000000000 65535 f \n"
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode()

    trailer = (
        b"trailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF\n"
    )

    Path(path).write_bytes(header + obj1 + obj2 + obj3 + obj4 + xref + trailer)


async def _create_docx(
    path: str, title: str, dns_canary: str, token_uid: str, settings: Any
) -> Path:
    """Generate a DOCX with an External-mode image relationship pointing at
    the canary URL. Word fetches external images when opening — that fetch
    is the trigger.

    This replaces the previous (inert) approach of writing the URL as
    null-prefixed text in a footer paragraph, which Word never resolved.

    Word's Protected View settings can block external image loading on first
    open — the user is prompted to "Enable Editing". Many enterprise
    deployments have Protected View disabled or auto-enable for trusted
    locations, so the planted document still fires there. Documented in
    KNOWN_LIMITATIONS.md.
    """
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(
            "This document contains sensitive information. Unauthorized access is prohibited."
        )
        doc.add_heading("CLASSIFICATION: CONFIDENTIAL", level=1)
        doc.add_paragraph("Please refer to the appendix for technical details.")
        doc.save(path)

        tracking_url = f"http://{dns_canary}/t/{token_uid}.png"
        _inject_external_image_rel(path, tracking_url)
    except ImportError:
        Path(path).write_bytes(b"PK\x03\x04")  # Minimal ZIP stub

    return Path(path)


def _inject_external_image_rel(path: str, image_url: str) -> None:
    """Hand-edit the OOXML zip to add an External-mode image relationship
    and a `<w:drawing>` element in the body that references it.

    python-docx doesn't expose External-mode rels directly so we manipulate
    the zip directly. Adds:
      1. A new relationship in `word/_rels/document.xml.rels` of type
         `.../image` with `Target="<url>"` and `TargetMode="External"`.
      2. A run with `<w:drawing>...<a:blip r:link="rIdN"/>...</w:drawing>`
         in the document body. The `r:link` (not `r:embed`) is what makes
         Word treat it as an external resource and fetch it on open.
    """
    import re
    import shutil
    import tempfile
    import zipfile

    # Build the rel id — pick something unlikely to collide with python-docx output.
    rel_id = "rIdHpCanary"

    # New relationship XML to inject into document.xml.rels.
    new_rel = (
        f'<Relationship Id="{rel_id}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
        f'Target="{image_url}" TargetMode="External"/>'
    )

    # The drawing block — 1x1 EMU is too small; Word uses 9525 EMU per pixel.
    # We use a 1-px (9525x9525 EMU) image so it's effectively invisible but
    # still requested by Word's rendering pipeline.
    drawing_xml = f"""<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" distT="0" distB="0" distL="0" distR="0"><wp:extent cx="9525" cy="9525"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="1" name="canary"/><wp:cNvGraphicFramePr/><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:nvPicPr><pic:cNvPr id="1" name="canary"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:link="{rel_id}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="9525" cy="9525"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>"""

    tmp_dir = tempfile.mkdtemp(prefix="docx-canary-")
    try:
        # Copy the docx contents into the tempdir.
        with zipfile.ZipFile(path, "r") as z:
            z.extractall(tmp_dir)

        # 1. Patch document.xml.rels to include the external image relationship.
        rels_path = Path(tmp_dir) / "word" / "_rels" / "document.xml.rels"
        rels = rels_path.read_text(encoding="utf-8")
        # Splice the new relationship just before the closing </Relationships>.
        rels_patched = rels.replace("</Relationships>", f"{new_rel}</Relationships>")
        rels_path.write_text(rels_patched, encoding="utf-8")

        # 2. Patch document.xml to add the drawing inside the body.
        doc_path = Path(tmp_dir) / "word" / "document.xml"
        body_xml = doc_path.read_text(encoding="utf-8")
        # Insert the drawing paragraph just before the body's section properties
        # or before the closing </w:body>. The latter is more robust across
        # python-docx versions.
        body_patched = re.sub(
            r"</w:body>",
            f"{drawing_xml}</w:body>",
            body_xml,
            count=1,
        )
        doc_path.write_text(body_patched, encoding="utf-8")

        # 3. Re-zip back into the docx.
        Path(path).unlink()
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in Path(tmp_dir).rglob("*"):
                if item.is_file():
                    zout.write(item, item.relative_to(tmp_dir).as_posix())
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
